import os
import re
import docx
from docx.shared import Inches, Pt
import pandas as pd
from Bio.Blast import NCBIWWW
from Bio.Blast import NCBIXML
from Bio import Entrez
Entrez.email = "fetchmysequencefromgenbank@gmail.com"

class compareSequencesNCBI:
    """
    This class will get all the sequences the user wants to identify and run a web blast against the genbank database
    It will also run a blast against the suposed species if given
    
    It is perfered if the input file is a docx file. a txt file is also accepted but the user should be careful with the simbols put in the file
    
    """

    def __init__(self, directory="",folderDirectoryToSaveFiles="", fileName ="", listOfDicts=[], fileExtension=""):
        self.directory = directory ##dictory of the folder where the initial input excel should be saved
        self.folderDirectoryToSaveFiles = folderDirectoryToSaveFiles #  #directory of the folder where the result file should be saved
        self.fileName = fileName #name of the input file with correct extension
        self.extensions = ['.doc', '.docm', '.docx', '.txt'] #list of acceptable extension files
        self.dnaRegex = re.compile(r"^([ACTGWSMKRYN])+$") # regex expression to retrieve only dna strings
        self.listOfDicts = listOfDicts #list of dictionaries where  header: __ nameOfSpecies : __, DNASeq: __, blastNameGeneral: __, blastNameEspecific  will be added from de input file
        self.fileExtension = fileExtension # extension of the file given by the user
        # directoria : path= r'C:\Users\jrrs1\OneDrive\Ambiente de Trabalho\Joana\bolsa\BI_IPB\Plantas\Informatica\compareSequencesNCBI'
        # nomes dos ficheiros : teste1-3.docx teste4-6.txt

    def getDirectoryFromUser(self):
        """
        Functions gets directory where the user has the input file and where they want to save the output files.
        The user must copy the directory of the folder(s) and paste it and hit ENTER.
        """
        
        if self.directory == "":    
            folderDirectory = input("Tell me which directory you have the input file: (when done click ENTER)")
            if os.path.exists(folderDirectory):
                self.directory = folderDirectory
            else:
                self.directory = os.getcwd()

            folderDirectoryToSaveFiles = input("Tell me which directory you want to save your files: (when done click ENTER)")
            if os.path.exists(folderDirectoryToSaveFiles):
                self.folderDirectoryToSaveFiles = folderDirectoryToSaveFiles
            else:
                self.folderDirectoryToSaveFiles = self.directory
            os.chdir(self.directory)


    def getFileNameFromUser(self):
        """
        Function gets the name of the input file where the user saved all the fasta sequences and headers they want to search.
        The user must write the name of the input file without the termination. 
        """
        if self.fileName == "":
            fileName = input("Tell me the name of the input file:  (when done click ENTER)")
            for extension in self.extensions:
                tempFileName= fileName + extension
                if os.path.exists(tempFileName):
                    self.fileName = tempFileName
                    self.fileExtension = extension
            if self.fileName=="":
                print("The input file needs to be saved in the same directory as given in the first step")
        else:
            self.fileName = self.fileName.split("/")[-1]
            self.fileExtension = self.fileName.split(".")[1]
            

    def processInputFile(self):
        count = 0
        if self.fileExtension == ".txt":
            with open(self.fileName,'r') as f:
                lines = f.readlines()
                for i in range(len(lines)):
                    temp = self.dnaRegex.search(lines[i])
                    if temp != None and temp.span()[1] >20:
                        first_name = re.search('([A-Z]([a-z])+)', lines[i-1])
                        tempList = lines[i-1].split(" ")
                        if len(tempList) > 1:
                            for index in range(len(tempList)):
                                if tempList[index] == first_name.group():
                                    temp_species_name = tempList[index] + " " + tempList[index +1]
                                    doc_name_general = "my_blast_" + "_".join(lines[i-1].strip().split(" ")) + "_general.xml"
                                    doc_name_specific = "my_blast_" + "_".join(lines[i-1].strip().split(" ")) + "_specific.xml"
                        else:
                            temp_species_name = "unknown" + str(count)
                            doc_name_general = "my_blast_" + "_".join(temp_species_name.split(" ")) + "_general.xml"
                            doc_name_specific = "my_blast_" + "_".join(temp_species_name.split(" ")) + "_specific.xml"
                            count+=1
                        tempDict = {'header': lines[i-1].strip(), 'species name': temp_species_name, 'DNASeq': lines[i].strip(), 'blastNameGeneral': doc_name_general, 'blastNameSpecific': doc_name_specific}
                        self.listOfDicts.append(tempDict)
        else:
            doc = docx.Document(self.fileName)
            for i in range(len(doc.paragraphs)):
                temp = self.dnaRegex.search(doc.paragraphs[i].text)
                if temp != None and temp.span()[1] > 20 :
                    first_name = re.search('([A-Z]([a-z])+)', doc.paragraphs[i-1].text)
                    tempList = doc.paragraphs[i-1].text.split(" ")
                    if len(tempList) > 1:
                        for index in range(len(tempList)):
                            if tempList[index] == first_name.group():
                                temp_species_name = tempList[index] + " " + tempList[index +1]
                                doc_name_general = "my_blast_" + "_".join(doc.paragraphs[i-1].text.split(" ")) + "_general.xml"
                                doc_name_specific = "my_blast_" + "_".join(doc.paragraphs[i-1].text.split(" ")) + "_specific.xml"
                    else:
                        temp_species_name = "unknown" + str(count)
                        doc_name_general = "my_blast_" + "_".join(temp_species_name.split(" ")) + "_general.xml"
                        doc_name_specific = "my_blast_" + "_".join(temp_species_name.split(" ")) + "_specific.xml"
                        count+=1
                    tempDict = {'header': doc.paragraphs[i-1].text, 'species name': temp_species_name, 'DNASeq': doc.paragraphs[i].text, 'blastNameGeneral': doc_name_general,  'blastNameSpecific': doc_name_specific}
                    self.listOfDicts.append(tempDict)

    def saveBLAST(self):
        os.chdir(self.folderDirectoryToSaveFiles)
        for dict in self.listOfDicts:
            keys=[dict['blastNameGeneral'], dict['blastNameSpecific']]
            for key in keys: #safe to continue the program where it left of if it is interropted
                tempPath = os.path.join(self.folderDirectoryToSaveFiles, key)
                if os.path.exists(tempPath):
                    continue
                print("im searching")
                result_handle_geral = NCBIWWW.qblast("blastn", "nt", dict['DNASeq'], megablast=True, hitlist_size = 20)
                result_handle_specific = NCBIWWW.qblast("blastn", "nt", dict['DNASeq'], megablast=True, entrez_query = dict['species name'] + "[ORGN]", hitlist_size = 10)
                with open(dict['blastNameGeneral'], "w") as f:
                    f.write(result_handle_geral.read())
                    result_handle_geral.close()
                f.close()
                with open(dict['blastNameSpecific'], "w") as f:
                    f.write(result_handle_specific.read())
                    result_handle_specific.close()
                f.close()
        #print("im done saving file")

    def processBLASTresults(self):
        for dict in self.listOfDicts:
            listOfRecordsNames = [dict['blastNameGeneral'], dict['blastNameSpecific']]
            for record in listOfRecordsNames: 
                df = pd.DataFrame()
                description, scientific_name, max_score, total_score, coverage, e_value, identity, acc_len, accession = [] ,[], [], [], [], [], [], [], []
                blast_record = NCBIXML.parse(open(record, 'r'))
                for query in blast_record:
                    for hit in query.alignments:
                        accession.append(hit.hit_id.split('|')[3])
                        acc_len.append(hit.length)
                        first = True
                        tempTotal_score = 0
                        for hsp in hit.hsps:
                            if first:
                                tempTotal_score += hsp.bits
                                max_score.append(hsp.bits)
                                description.append(hit.hit_def)
                                scientific_name.append(" ".join(hit.hit_def.split(" ")[0:2]))
                                coverage.append(round((hsp.query_end - hsp.query_start +1 ) / query.query_length * 100, 2))
                                identity.append(round((hsp.identities/ hsp.align_length * 100), 2))
                                e_value.append(hsp.expect)
                                first = False
                            else:
                                tempTotal_score += hsp.bits
                        total_score.append(tempTotal_score)
                        tempTotal_score=0
                        first=True
                df_name = record.split("_")[-1] + "_df"
                df['Description'],  df['Scientific Name'], df['Max Score'], df['Total Score'], df['Query Cover'], df['E value'], df['Per. Ident'], df['Acc. Len'], df['Accession'] = description, scientific_name, max_score, total_score, coverage, e_value, identity, acc_len, accession
                dict[df_name]=df

    def constructResultsFile(self):
        
        resultFileName = "results_" + self.fileName.split(".")[0] + ".docx"
        outputFile = docx.Document()
        sections = outputFile.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.right_margin  = Inches(0.75)
            section.left_margin = Inches(0.75)
        font = outputFile.styles['Normal'].font
        font.name = 'Arial'
        font.size = Pt(8)
        for dict in self.listOfDicts:
            outputFile.add_paragraph(dict['header'])
            listOfDf = ['general.xml_df', 'specific.xml_df']
            for df in listOfDf:
                outputFile.add_paragraph(df.split(".")[0])
                t = outputFile.add_table(dict[df].shape[0]+1, dict[df].shape[1]) # extra row is so we can add the header row
                t.style = 'Light Grid Accent 5'
                for j in range(dict[df].shape[-1]): #add the header
                    t.cell(0,j).text = dict[df].columns[j]
                for i in range(dict[df].shape[0]): # add the rest of the data frame
                    for j in range(dict[df].shape[-1]):
                        t.cell(i+1,j).text = str(dict[df].values[i,j])
                listOfDimensions = [2.32, 1, 0.56, 0.56, 0.49, 0.45, 0.43, 0.54, 0.84]
                for column in range(len(t.columns)):
                    for cell in t.columns[column].cells:
                        cell.width = Inches(listOfDimensions[column])
        outputFile.save(resultFileName)
        
    def removeIntermediateFiles(self):
        for dict in self.listOfDicts:
            keys = ['blastNameGeneral', 'blastNameSpecific']
            for key in keys:
                tempDirectory = os.path.join(self.folderDirectoryToSaveFiles, dict[key])
                if os.path.exists(tempDirectory):
                    os.remove(tempDirectory)

#if __name__ == "__main__":
    #teste = compareSequencesNCBI()
    #teste.getDirectoryFromUser()
    #teste.getFileNameFromUser()
    #teste.processInputFile()
    #teste.saveBLAST()
    #teste.processBLASTresults()
    #teste.constructResultsFile()
    #teste.removeIntermediateFiles()
# -*- coding: utf-8 -*-
"""
Created on Tue Mar  7 15:44:30 2023

@author: Joana Santos
"""
import docx
import re
import pandas as pd


path=r'C:\Users\jrrs1\OneDrive\Ambiente de Trabalho\Joana\bolsa\BI_IPB\Plantas\Informatica\getSpeciesFromPDF\Farmacopeia Europeia 8-1061-1352.docx'
path2= r'C:\Users\jrrs1\OneDrive\Ambiente de Trabalho\Joana\bolsa\BI_IPB\Plantas\Informatica\getSpeciesFromPDF\results_teste.docx'



doc = docx.Document(path)
all_paras = doc.paragraphs
print(len(all_paras))

mydoc = docx.Document()

for i in range(len(all_paras)):
    if all_paras[i].text == "DEFINITION":
        mydoc.add_paragraph(all_paras[i].text)
        mydoc.add_paragraph(all_paras[i+1].text)
        
mydoc.save(path2)
    
    #print(para.text)
    #print("-------")
    


doc = docx.Document(path2)


count=0
with open('species_control.txt', 'w', encoding="utf-8") as f:
    countAfterProcess=0    
    for para in doc.paragraphs:
        splited = para.text.split()
        for i in range(len(splited)):
            if splited[i]=="of":
                count += 1
                if len(splited[i:])>2 and re.findall('([A-Z][a-z]+)', splited[i+1]):
                    countAfterProcess += 1
                    f.write(" ".join(splited[i:]))
                    f.write('\n')
                    f.write(" ".join(splited[i+1:i+3]))
                    f.write('\n')
    print(count)
    print(countAfterProcess)

species=[]
df = pd.read_excel("Medicines_output_herbal_medicines_Species_names.xlsx")

species = df["Especies"].tolist()
with open("species_control2.txt" ,"w", encoding="utf-8") as f:
          for specie in species:
              f.write(specie)
              f.write('\n')

for i in range(len(species)):
    species[i]= " ".join(species[i].split()[0:2])
    print(species[i])

with open('species2.txt', 'a',  encoding="utf-8") as f:
    f.write('\n')
    for specie in species:
        f.write(specie)
        f.write('\n')